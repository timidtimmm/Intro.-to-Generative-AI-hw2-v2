"""
rag.py — Retrieval-Augmented Generation (RAG) Module
======================================================
Supports:
  - PDF / TXT / MD / DOCX document ingestion   ★ 新增 .docx
  - Web URL scraping                            ★ 新增 URL 抓取
  - Text chunking with overlap
  - Embedding via NVIDIA NIM API (or TF-IDF fallback)
  - SQLite vector store (cosine similarity)
  - Context retrieval for RAG-enhanced chat

Usage:
  from rag import rag_manager
  doc_id = rag_manager.ingest(file_bytes, filename, session_id)
  context = rag_manager.retrieve(query, session_id, top_k=4)
"""

import hashlib
import json
import math
import os
import re
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Tuple

import httpx

DB_PATH = Path("chat_history.db")
NVIDIA_API_KEY = os.getenv("NVIDIA_API_KEY", "")
NVIDIA_BASE_URL = os.getenv("NVIDIA_BASE_URL", "https://integrate.api.nvidia.com/v1")
EMBED_MODEL = os.getenv("NVIDIA_EMBED_MODEL", "nvidia/nv-embedqa-e5-v5")

CHUNK_SIZE = 500        # characters per chunk
CHUNK_OVERLAP = 80      # overlap between chunks
MAX_CHUNKS_PER_DOC = 300


# ── DB Init ───────────────────────────────────────────────────────────────────

def init_rag_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS rag_documents (
            id          TEXT PRIMARY KEY,
            session_id  TEXT,
            filename    TEXT,
            content_hash TEXT,
            chunk_count INTEGER,
            created_at  TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS rag_chunks (
            id          TEXT PRIMARY KEY,
            doc_id      TEXT,
            session_id  TEXT,
            chunk_idx   INTEGER,
            content     TEXT,
            embedding   TEXT,   -- JSON array of floats
            created_at  TEXT,
            FOREIGN KEY (doc_id) REFERENCES rag_documents(id)
        )
    """)
    conn.commit()
    conn.close()


init_rag_db()


# ── Text Extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(data: bytes) -> str:
    # ★ 優先用 pymupdf（對壓縮字型支援最好）
    try:
        import io, fitz  # pymupdf
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for i, page in enumerate(doc, start=1):
            if i > 50:
                break
            text = page.get_text("text")
            if text.strip():
                pages.append(f"[Page {i}]\n{text}")
        doc.close()
        result = "\n\n".join(pages)
        if len(result.strip()) > 50:
            return result
    except Exception as e:
        print(f"[RAG] pymupdf error: {e}")

    # 次選：pdfplumber
    try:
        import io, pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages[:50], start=1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i}]\n{text}")
            result = "\n\n".join(pages)
            if len(result.strip()) > 50:
                return result
    except Exception as e:
        print(f"[RAG] pdfplumber error: {e}")

    # 最後 fallback：pypdf
    try:
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for i, page in enumerate(reader.pages[:50], start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {i}]\n{text}")
        return "\n\n".join(pages)
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def extract_text_from_docx(data: bytes) -> str:
    """
    ★ 新增：從 .docx 二進位資料提取純文字。
    優先用 mammoth（保留段落結構），fallback 到 python-docx，
    再 fallback 到錯誤訊息。
    安裝: pip install mammoth python-docx --break-system-packages
    """
    import io

    # 方法一：mammoth（忽略樣式，只取文字，最乾淨）
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(data))
        text = result.value or ""
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"[RAG] mammoth error: {e}")

    # 方法二：python-docx（同時抓段落 + 表格）
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts: list[str] = []

        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        text = "\n\n".join(parts)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"[RAG] python-docx error: {e}")

    return "[DOCX extraction failed: install mammoth or python-docx]"


def extract_text_from_url(url: str, timeout: float = 15.0) -> Tuple[str, str]:
    """
    ★ 新增：抓取網頁並提取正文。
    回傳 (text, resolved_url)。
    安裝: pip install beautifulsoup4 --break-system-packages
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError(
            "beautifulsoup4 not installed. Run: pip install beautifulsoup4 --break-system-packages"
        )

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (compatible; RAGBot/1.0)"
        )
    }

    resp = httpx.get(url, headers=headers, timeout=timeout, follow_redirects=True)
    resp.raise_for_status()
    resolved_url = str(resp.url)
    content_type = resp.headers.get("content-type", "")

    # 純文字頁面直接回傳
    if "text/html" not in content_type and "application/xhtml" not in content_type:
        return resp.text[:40_000], resolved_url

    soup = BeautifulSoup(resp.text, "html.parser")

    # 移除雜訊標籤
    for tag in soup(["script", "style", "nav", "footer", "header",
                     "aside", "iframe", "noscript", "form", "button"]):
        tag.decompose()

    # 優先取語意容器
    main = (
        soup.find("main")
        or soup.find("article")
        or soup.find(id=re.compile(r"content|main|article", re.I))
        or soup.find(class_=re.compile(r"content|main|article|post-body", re.I))
        or soup.body
        or soup
    )

    lines: list[str] = []
    for elem in main.find_all(
        ["h1", "h2", "h3", "h4", "h5", "h6",
         "p", "li", "td", "th", "pre", "code", "blockquote"]
    ):
        text = elem.get_text(separator=" ", strip=True)
        if text and len(text) > 20:
            lines.append(text)

    text = "\n\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text[:40_000], resolved_url


def _is_docx(filename: str, mime: str) -> bool:
    return (
        (filename or "").lower().endswith(".docx")
        or "wordprocessingml" in mime
        or mime == "application/msword"
    )


def _is_url(text: str) -> bool:
    """檢查字串是否為 http/https URL。"""
    return bool(re.match(r"^https?://\S+", (text or "").strip()))


def extract_text(data: bytes, filename: str, mime: str = "") -> str:
    """統一入口：依格式分派到對應的 extractor。"""
    fname = (filename or "").lower()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if mime == "application/pdf" or fname.endswith(".pdf"):
        return extract_text_from_pdf(data)

    # ── DOCX ★ 新增 ───────────────────────────────────────────────────────────
    if _is_docx(filename, mime):
        return extract_text_from_docx(data)

    # ── URL ★ 新增（data 為空，filename 傳入 URL）─────────────────────────────
    if _is_url(filename):
        try:
            text, _ = extract_text_from_url(filename)
            return text
        except Exception as e:
            return f"[URL fetch error: {e}]"

    # ── 純文字 ────────────────────────────────────────────────────────────────
    try:
        return data.decode("utf-8", errors="replace")
    except Exception:
        return ""


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks, preferring sentence boundaries."""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            for sep in ["\n\n", "。", ".\n", ". ", "\n"]:
                pos = text.rfind(sep, start + chunk_size // 2, end)
                if pos != -1:
                    end = pos + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap

    return chunks[:MAX_CHUNKS_PER_DOC]


# ── Embedding ─────────────────────────────────────────────────────────────────

def _embed_nvidia(texts: List[str]) -> Optional[List[List[float]]]:
    if not NVIDIA_API_KEY or not texts:
        return None
    try:
        resp = httpx.post(
            f"{NVIDIA_BASE_URL}/embeddings",
            headers={"Authorization": f"Bearer {NVIDIA_API_KEY}", "Content-Type": "application/json"},
            json={"model": EMBED_MODEL, "input": texts, "input_type": "passage", "encoding_format": "float"},
            timeout=30.0,
        )
        if resp.status_code >= 400:
            print(f"[RAG] Embedding API error {resp.status_code}: {resp.text[:300]}")
            return None
        data = resp.json()
        embeddings = [item["embedding"] for item in sorted(data["data"], key=lambda x: x["index"])]
        return embeddings
    except Exception as e:
        print(f"[RAG] Embedding error: {e}")
        return None


def _tfidf_embed(text: str) -> List[float]:
    tokens = re.findall(r"[\u4e00-\u9fff]|[a-zA-Z0-9]+", text.lower())
    freq: Dict[str, int] = {}
    for t in tokens:
        freq[t] = freq.get(t, 0) + 1
    if not freq:
        return [0.0] * 128
    vec = [0.0] * 128
    for word, count in freq.items():
        idx = int(hashlib.md5(word.encode()).hexdigest(), 16) % 128
        vec[idx] += count / len(tokens)
    norm = math.sqrt(sum(x*x for x in vec)) or 1.0
    return [x / norm for x in vec]


def embed_texts(texts: List[str]) -> List[List[float]]:
    nvidia_result = _embed_nvidia(texts)
    if nvidia_result and len(nvidia_result) == len(texts):
        return nvidia_result
    print("[RAG] Using TF-IDF fallback embedding")
    return [_tfidf_embed(t) for t in texts]


def embed_query(query: str) -> List[float]:
    if NVIDIA_API_KEY:
        try:
            resp = httpx.post(
                f"{NVIDIA_BASE_URL}/embeddings",
                headers={"Authorization": f"Bearer {NVIDIA_API_KEY}", "Content-Type": "application/json"},
                json={"model": EMBED_MODEL, "input": [query], "input_type": "query", "encoding_format": "float"},
                timeout=15.0,
            )
            if resp.status_code < 400:
                return resp.json()["data"][0]["embedding"]
        except Exception as e:
            print(f"[RAG] Query embedding error: {e}")
    return _tfidf_embed(query)


# ── Similarity ────────────────────────────────────────────────────────────────

def cosine_similarity(a: List[float], b: List[float]) -> float:
    if len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x*x for x in a))
    nb = math.sqrt(sum(x*x for x in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


# ── RAG Manager ───────────────────────────────────────────────────────────────

class RAGManager:

    def ingest(self, data: bytes, filename: str, session_id: str, mime: str = "") -> Dict:
        """
        Ingest a document: extract text → chunk → embed → store in DB.
        ★ 新增支援 .docx 和 URL（filename 傳入 URL 字串，data 傳空 bytes）。
        Returns {"doc_id": ..., "chunk_count": ..., "filename": ...}
        """
        # ── URL 模式 ★ 新增 ──────────────────────────────────────────────────
        if _is_url(filename):
            try:
                text, resolved_url = extract_text_from_url(filename)
            except Exception as e:
                return {"error": f"URL 抓取失敗：{e}"}

            if not text.strip():
                return {"error": "網頁沒有可讀內容"}

            # 用 URL 內容的 hash 做去重（不用 data hash）
            content_hash = hashlib.sha256(text.encode()).hexdigest()
            display_name = resolved_url  # 顯示完整 URL

            conn = sqlite3.connect(DB_PATH)
            existing = conn.execute(
                "SELECT id, chunk_count FROM rag_documents WHERE session_id=? AND content_hash=?",
                (session_id, content_hash)
            ).fetchone()
            conn.close()
            if existing:
                return {"doc_id": existing[0], "chunk_count": existing[1],
                        "filename": display_name, "status": "already_ingested"}

            return self._store_text(text, display_name, session_id, content_hash)

        # ── 一般檔案模式 ─────────────────────────────────────────────────────
        content_hash = hashlib.sha256(data).hexdigest()

        conn = sqlite3.connect(DB_PATH)
        existing = conn.execute(
            "SELECT id, chunk_count FROM rag_documents WHERE session_id=? AND content_hash=?",
            (session_id, content_hash)
        ).fetchone()
        conn.close()
        if existing:
            return {"doc_id": existing[0], "chunk_count": existing[1],
                    "filename": filename, "status": "already_ingested"}

        text = extract_text(data, filename, mime)
        if not text.strip():
            return {"error": "無法從文件提取文字"}

        return self._store_text(text, filename, session_id, content_hash)

    def _store_text(self, text: str, filename: str, session_id: str, content_hash: str) -> Dict:
        """Chunk → embed → write to SQLite."""
        chunks = chunk_text(text)
        if not chunks:
            return {"error": "文件沒有產生任何片段"}

        # 批次 embed
        BATCH = 32
        all_embeddings: list = []
        for i in range(0, len(chunks), BATCH):
            batch = chunks[i:i + BATCH]
            embs = embed_texts(batch)
            all_embeddings.extend(embs)

        doc_id = str(uuid.uuid4())
        now = datetime.utcnow().isoformat()

        conn = sqlite3.connect(DB_PATH)
        conn.execute(
            "INSERT INTO rag_documents (id, session_id, filename, content_hash, chunk_count, created_at) VALUES (?,?,?,?,?,?)",
            (doc_id, session_id, filename, content_hash, len(chunks), now)
        )
        for idx, (chunk, emb) in enumerate(zip(chunks, all_embeddings)):
            conn.execute(
                "INSERT INTO rag_chunks (id, doc_id, session_id, chunk_idx, content, embedding, created_at) VALUES (?,?,?,?,?,?,?)",
                (str(uuid.uuid4()), doc_id, session_id, idx, chunk, json.dumps(emb), now)
            )
        conn.commit()
        conn.close()

        return {"doc_id": doc_id, "chunk_count": len(chunks),
                "filename": filename, "status": "ingested"}

    def retrieve(self, query: str, session_id: str, top_k: int = 4, min_score: float = 0.15) -> List[Dict]:
        """
        Retrieve the most relevant chunks for a query within a session.
        Returns list of {"content": ..., "score": ..., "filename": ..., "chunk_idx": ...}
        """
        conn = sqlite3.connect(DB_PATH)
        rows = conn.execute(
            """SELECT rc.content, rc.embedding, rd.filename, rc.chunk_idx
               FROM rag_chunks rc
               JOIN rag_documents rd ON rd.id = rc.doc_id
               WHERE rc.session_id = ?
               ORDER BY rc.doc_id, rc.chunk_idx""",
            (session_id,)
        ).fetchall()
        conn.close()

        if not rows:
            return []

        q_emb = embed_query(query)

        scored = []
        for content, emb_json, filename, chunk_idx in rows:
            try:
                emb = json.loads(emb_json)
            except Exception:
                continue
            score = cosine_similarity(q_emb, emb)
            if score >= min_score:
                scored.append({
                    "content": content,
                    "score": round(score, 4),
                    "filename": filename,
                    "chunk_idx": chunk_idx,
                })

        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[:top_k]

    def list_docs(self, session_id: str) -> List[Dict]:
        conn = sqlite3.connect(DB_PATH)
        rows = conn.execute(
            "SELECT id, filename, chunk_count, created_at FROM rag_documents WHERE session_id=? ORDER BY created_at DESC",
            (session_id,)
        ).fetchall()
        conn.close()
        return [{"id": r[0], "filename": r[1], "chunk_count": r[2], "created_at": r[3]} for r in rows]

    def delete_doc(self, doc_id: str):
        conn = sqlite3.connect(DB_PATH)
        conn.execute("DELETE FROM rag_chunks WHERE doc_id=?", (doc_id,))
        conn.execute("DELETE FROM rag_documents WHERE id=?", (doc_id,))
        conn.commit()
        conn.close()

    def format_context(self, chunks: List[Dict], max_chars: int = 3000) -> str:
        if not chunks:
            return ""
        lines = ["[RAG 檢索到的相關文件內容 — 請優先參考以下資訊回答使用者問題:]"]
        total = 0
        for i, c in enumerate(chunks, 1):
            entry = f"\n--- 來源：{c['filename']} (片段 {c['chunk_idx']+1}, 相似度 {c['score']:.2f}) ---\n{c['content']}"
            if total + len(entry) > max_chars:
                break
            lines.append(entry)
            total += len(entry)
        lines.append("\n[以上為檢索結果，請根據這些內容回答，並在回答中標注資訊來源。]")
        return "\n".join(lines)


rag_manager = RAGManager()